import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from datetime import datetime, date
import io
import re

# --- 1. CONFIGURAÇÃO VISUAL ---
st.set_page_config(page_title="Gerador PCPE Oficial", layout="wide", page_icon="🚓")

st.markdown("""
    <style>
    .main {background-color: #f8f9fa;}
    .stTextInput>div>div>input {font-weight: 500; color: #000;}
    .stTextArea textarea {font-family: 'Arial'; font-size: 14px;}
    .ferramenta-box {
        background-color: #fff3cd; padding: 10px; border-radius: 5px; 
        border: 1px solid #ffeeba; margin-bottom: 10px; font-size: 13px;
    }
    </style>
""", unsafe_allow_html=True)

# --- 2. FUNÇÕES DE INTELIGÊNCIA INVESTIGATIVA (METADADOS) ---

def converter_para_graus(valor):
    """Converte coordenadas GPS (DMS) para Decimal."""
    d = float(valor[0])
    m = float(valor[1])
    s = float(valor[2])
    return d + (m / 60.0) + (s / 3600.0)

def get_image_metadata(image_file):
    """Extrai Data e GPS da imagem."""
    info_dict = {"data": None, "gps": None}
    
    try:
        image = Image.open(image_file)
        exif_data = image._getexif()
        
        if not exif_data:
            return info_dict

        # 1. Extração da Data
        for tag, value in exif_data.items():
            tag_name = TAGS.get(tag, tag)
            if tag_name == 'DateTimeOriginal':
                try:
                    dt_obj = datetime.strptime(value, '%Y:%m:%d %H:%M:%S')
                    info_dict["data"] = dt_obj.strftime('%d/%m/%Y às %Hh%M')
                except:
                    pass
        
        # 2. Extração do GPS
        if 'GPSInfo' in [TAGS.get(x, x) for x in exif_data]:
            gps_info = {}
            for key in exif_data.keys():
                decode = TAGS.get(key, key)
                if decode == "GPSInfo":
                    for t in exif_data[key]:
                        sub_decode = GPSTAGS.get(t, t)
                        gps_info[sub_decode] = exif_data[key][t]
            
            if 'GPSLatitude' in gps_info and 'GPSLongitude' in gps_info:
                lat = converter_para_graus(gps_info['GPSLatitude'])
                lon = converter_para_graus(gps_info['GPSLongitude'])
                
                # Ajuste de Hemisfério (Sul/Oeste são negativos)
                if gps_info.get('GPSLatitudeRef') == 'S': lat = -lat
                if gps_info.get('GPSLongitudeRef') == 'W': lon = -lon
                
                info_dict["gps"] = f"{lat:.5f}, {lon:.5f}"

    except Exception:
        pass # Se der erro na leitura, retorna nulo sem quebrar o app
        
    return info_dict

def calcular_idade(nascimento):
    """Calcula idade exata."""
    today = date.today()
    idade = today.year - nascimento.year - ((today.month, today.day) < (nascimento.month, nascimento.day))
    return idade

# --- 3. FUNÇÕES DE FORMATAÇÃO (MANTIDAS) ---
def formatar_texto(run, tamanho=11, negrito=False, italico=False):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(tamanho)
    run.bold = negrito
    run.italic = italico
    run.font.color.rgb = RGBColor(0, 0, 0)

def configurar_paragrafo(paragrafo, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, espaco_depois=0, entrelinhas=1.0, recuo=0):
    p_fmt = paragrafo.paragraph_format
    p_fmt.alignment = alinhamento
    p_fmt.space_after = Pt(espaco_depois)
    p_fmt.line_spacing = entrelinhas
    if recuo > 0: p_fmt.first_line_indent = Cm(recuo)

# --- 4. CABEÇALHO PERFEITO (MANTIDO) ---
def criar_cabecalho_rodape(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.4); section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.2); section.footer_distance = Inches(0.2)

    header = section.header
    table = header.add_table(rows=1, cols=3, width=Inches(7.2))
    table.autofit = False
    
    largura_lateral = Inches(1.3)
    largura_central = Inches(4.6)
    table.columns[0].width = largura_lateral
    table.columns[1].width = largura_central
    table.columns[2].width = largura_lateral

    try:
        cell_logo = table.cell(0, 0)
        cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        run_logo.add_picture('logo_pc.png', width=Inches(1.0))
    except:
        table.cell(0, 0).text = "[LOGO]"

    cell_text = table.cell(0, 1)
    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_text._element.clear_content()

    def add_line(texto, tamanho):
        p = cell_text.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(texto)
        formatar_texto(r, tamanho=tamanho, negrito=True)

    add_line("POLÍCIA CIVIL DE PERNAMBUCO", 14)
    add_line("DINTER 1 - 16ª DESEC", 11)
    add_line("Delegacia de Polícia da 116ª Circunscrição - Surubim", 11)

    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("Av. São Sebastião - Surubim - PE | Fone: (81) 3624-1974\nE-mail: dp116circ.surubim@policiacivil.pe.gov.br")
    formatar_texto(r_foot, tamanho=9)

# --- 5. INTERFACE ---
if 'num_agentes' not in st.session_state: st.session_state.num_agentes = 1
def add_agente(): st.session_state.num_agentes += 1
def remove_agente(): 
    if st.session_state.num_agentes > 1: st.session_state.num_agentes -= 1

# BARRA LATERAL
with st.sidebar:
    st.header("1. Cabeçalho")
    titulo_doc = st.text_input("Título:", value="RELATÓRIO DE INVESTIGAÇÃO")
    st.markdown("---")
    opj = st.text_input("OPJ:", placeholder="Ex: INTERCEPTUM")
    processo = st.text_input("Processo:", placeholder="0002343...")
    natureza = st.text_input("Natureza:", placeholder="Homicídio...")
    c1, c2 = st.columns(2)
    data_input = c1.text_input("Data:", placeholder="DD/MM/AAAA")
    hora_input = c2.text_input("Hora:", placeholder="HH:MM")
    local = st.text_input("Local:", placeholder="Endereço...")
    
    st.markdown("---")
    st.subheader("🧮 Calc. Idade Penal")
    dt_nasc_calc = st.date_input("Nascimento do Alvo:", value=None)
    if dt_nasc_calc:
        idade = calcular_idade(dt_nasc_calc)
        if idade < 18:
            st.error(f"⚠️ MENOR DE IDADE: {idade} anos")
        else:
            st.success(f"✅ IMPUTÁVEL: {idade} anos")

st.title("🚓 Gerador PCPE (Metadados Avançados)")

# ABAS
tab1, tab2, tab3 = st.tabs(["📝 Relato e Fotos", "👤 Envolvidos e Apreensão", "👮 Equipe"])

# Variável global
fotos_carregadas = []

with tab1:
    col_upload, col_texto = st.columns([1, 2])
    
    with col_upload:
        st.info("📷 Fotos (Data e GPS Automáticos)")
        fotos_carregadas = st.file_uploader("Upload", accept_multiple_files=True, label_visibility="collapsed")
        
        if fotos_carregadas:
            st.write("---")
            for i, f in enumerate(fotos_carregadas):
                with st.container():
                    c_img, c_code = st.columns([1, 2])
                    c_img.image(f, width=60)
                    
                    # Extração de Metadados
                    meta = get_image_metadata(f)
                    
                    # Construção da info para exibir
                    info_display = []
                    if meta['data']: info_display.append(f"📅 {meta['data']}")
                    if meta['gps']: info_display.append(f"📍 {meta['gps']}")
                    
                    c_code.code(f"[FOTO{i+1}]", language="html")
                    if info_display:
                        c_code.caption(" | ".join(info_display))

    with col_texto:
        st.subheader("Redação do Fato")
        texto_relato = st.text_area("Corpo do Texto:", height=600, 
                                   placeholder="Digite o histórico...\n\n[FOTO1]\n\n...")

with tab2:
    st.subheader("1. Qualificação")
    c_a, c_b = st.columns(2)
    with c_a:
        alvo = st.text_input("Nome Alvo:")
        cpf_rg = st.text_input("Docs (CPF/RG):")
        nasc = st.text_input("Nascimento:")
    with c_b:
        vitima = st.text_input("Nome Vítima:")
        advogado = st.text_input("Advogado:")
        testemunha = st.text_input("Testemunha:")
    
    st.markdown("---")
    st.subheader("2. Gerador de Lista de Apreensão 🔫")
    col_obj, col_qtd, col_desc = st.columns([2, 1, 3])
    obj = col_obj.text_input("Objeto (ex: Celular)")
    qtd = col_qtd.text_input("Qtd", value="1")
    desc = col_desc.text_input("Descrição (ex: Samsung, preto, IMEI...)")
    
    if 'lista_apreensao' not in st.session_state: st.session_state.lista_apreensao = []
    
    if st.button("⬇️ Adicionar à Lista"):
        if obj:
            texto_item = f"{qtd} (uma/ns) {obj}, {desc}"
            st.session_state.lista_apreensao.append(texto_item)
            st.success("Adicionado!")
    
    if st.session_state.lista_apreensao:
        st.markdown("**Itens adicionados:**")
        st.info("; ".join(st.session_state.lista_apreensao) + ".")
        if st.button("Limpar Lista"):
            st.session_state.lista_apreensao = []

with tab3:
    st.subheader("Assinaturas")
    agentes = []
    for i in range(st.session_state.num_agentes):
        c1, c2 = st.columns([3, 2])
        n = c1.text_input(f"Nome {i+1}", key=f"n{i}")
        c = c2.text_input(f"Cargo {i+1}", key=f"c{i}", value="Agente de Polícia")
        agentes.append((n, c))
    st.button("➕ Adicionar", on_click=add_agente)
    st.button("➖ Remover", on_click=remove_agente)

# --- 6. GERAÇÃO ---
st.markdown("---")
if st.button("GERAR RELATÓRIO FINAL", type="primary"):
    doc = Document()
    criar_cabecalho_rodape(doc)
    
    p_tit = doc.add_paragraph()
    r_tit = p_tit.add_run(titulo_doc.upper())
    formatar_texto(r_tit, tamanho=12, negrito=True)
    configurar_paragrafo(p_tit, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_depois=12)

    def add_dado(chave, valor):
        if valor:
            p = doc.add_paragraph()
            r_k = p.add_run(f"{chave}: ")
            formatar_texto(r_k, negrito=True)
            r_v = p.add_run(str(valor))
            formatar_texto(r_v, negrito=False)
            configurar_paragrafo(p, espaco_depois=0)

    add_dado("NATUREZA", natureza)
    add_dado("OPERAÇÃO (OPJ)", f"\"{opj}\"" if opj else None)
    add_dado("PROCESSO/BO", processo)
    if data_input and hora_input:
        add_dado("DATA/HORA", f"{data_input} às {hora_input}")
    elif data_input:
        add_dado("DATA", data_input)
    add_dado("LOCAL", local)
    
    doc.add_paragraph()

    if any([alvo, vitima, advogado, testemunha]):
        p_sec1 = doc.add_paragraph()
        r_sec1 = p_sec1.add_run("DOS ENVOLVIDOS")
        formatar_texto(r_sec1, negrito=True)
        configurar_paragrafo(p_sec1, espaco_depois=6)
        if alvo:
            txt = alvo
            if cpf_rg: txt += f" | {cpf_rg}"
            add_dado("ALVO/INVESTIGADO", txt)
            if nasc: add_dado("NASCIMENTO", nasc)
        add_dado("VÍTIMA", vitima)
        add_dado("ADVOGADO", advogado)
        add_dado("TESTEMUNHA", testemunha)
        doc.add_paragraph()
    
    if st.session_state.lista_apreensao:
        p_apr = doc.add_paragraph()
        r_apr = p_apr.add_run("DA APREENSÃO")
        formatar_texto(r_apr, negrito=True)
        configurar_paragrafo(p_apr, espaco_depois=6)
        
        texto_apreensao = "Durante a diligência, foram arrecadados os seguintes materiais: " + "; ".join(st.session_state.lista_apreensao) + "."
        p_itens = doc.add_paragraph(texto_apreensao)
        configurar_paragrafo(p_itens, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, entrelinhas=1.5, espaco_depois=6, recuo=1.25)
        for run in p_itens.runs: formatar_texto(run, tamanho=11)
        doc.add_paragraph()

    p_sec2 = doc.add_paragraph()
    r_sec2 = p_sec2.add_run("DO RELATO / DILIGÊNCIA")
    formatar_texto(r_sec2, negrito=True)
    configurar_paragrafo(p_sec2, espaco_depois=6)

    if texto_relato:
        partes = re.split(r'\[FOTO(\d+)\]', texto_relato)
        for parte in partes:
            if parte.isdigit():
                idx = int(parte) - 1
                if 0 <= idx < len(fotos_carregadas):
                    f = fotos_carregadas[idx]
                    
                    # Pega Metadados para a Legenda
                    meta = get_image_metadata(f)
                    partes_legenda = [f"Figura {idx+1}"]
                    if meta['data']: partes_legenda.append(f"Registro: {meta['data']}")
                    if meta['gps']: partes_legenda.append(f"Loc: {meta['gps']}")
                    texto_legenda = " | ".join(partes_legenda)
                    
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(f, width=Inches(5.5))
                    p_leg = doc.add_paragraph()
                    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_leg = p_leg.add_run(texto_legenda)
                    formatar_texto(r_leg, tamanho=9)
                    configurar_paragrafo(p_leg, espaco_depois=12)
            else:
                for par in parte.split('\n'):
                    if par.strip():
                        p = doc.add_paragraph(par)
                        configurar_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, entrelinhas=1.5, espaco_depois=6, recuo=1.25)
                        for run in p.runs: formatar_texto(run, tamanho=11)

    doc.add_paragraph(); doc.add_paragraph()
    for nome, cargo in agentes:
        if nome:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"__________________________________________\n{nome}\n{cargo}")
            formatar_texto(r, tamanho=11)

    bio = io.BytesIO()
    doc.save(bio)
    st.balloons()
    st.download_button("⬇️ BAIXAR DOCX", bio.getvalue(), "Relatorio_PCPE_Final_V25.docx", type="primary")
